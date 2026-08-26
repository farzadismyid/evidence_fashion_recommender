"""Rebuild thesis DOCX chapters from canonical Markdown sources."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
THESIS = ROOT / "thesis"


def clear_body(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag.endswith("}sectPr"):
            continue
        body.remove(child)


def latex_to_linear(source: str) -> str:
    """Convert the project’s limited LaTex vocabulary to readable linear math."""
    equation = source.strip()
    if equation.startswith("\\(") and equation.endswith("\\)"):
        equation = equation[2:-2]
    equation = re.sub(r"\\begin\{aligned\}|\\end\{aligned\}", "", equation)
    equation = equation.replace("\\\\", "\n")
    equation = re.sub(r"\\tag\{([^}]+)\}", r"  (\1)", equation)
    equation = re.sub(r"\\operatorname\{([^}]+)\}", r"\1", equation)
    equation = re.sub(r"\\mathrm\{([^}]+)\}", r"\1", equation)
    equation = equation.replace("\\mathbb{1}", "𝟙")
    equation = equation.replace("\\sum", "Σ").replace("\\min", "min").replace("\\max", "max")
    equation = equation.replace("\\times", "×").replace("\\top", "ᵀ")
    equation = equation.replace("\\lVert", "||").replace("\\rVert", "||")
    equation = equation.replace("\\dfrac", "\\frac")
    fraction = re.compile(r"\\frac\{([^{}]+)\}\{([^{}]+)\}")
    while fraction.search(equation):
        equation = fraction.sub(r"(\1)/(\2)", equation)
    equation = equation.replace("\\left", "").replace("\\right", "")
    equation = equation.replace("\\!", "")
    for latex, symbol in {
        "\\alpha": "α",
        "\\beta": "β",
        "\\gamma": "γ",
        "\\Delta": "Δ",
        "\\star": "⋆",
        "\\in": "∈",
        "\\neq": "≠",
        "\\ge": "≥",
        "\\le": "≤",
        "\\land": "∧",
        "\\varnothing": "∅",
        "\\models": "⊨",
    }.items():
        equation = equation.replace(latex, symbol)
    equation = equation.replace("\\", "").replace("{", "").replace("}", "")
    equation = re.sub(r"\^([A-Za-z0-9⋆]+)", r"⁽\1⁾", equation)
    return re.sub(r"_([A-Za-z0-9]+)", r"₍\1₎", equation)


def add_native_math(paragraph, source: str) -> None:
    math = OxmlElement("m:oMath")
    run = OxmlElement("m:r")
    text = OxmlElement("m:t")
    text.text = latex_to_linear(source)
    run.append(text)
    math.append(run)
    paragraph._p.append(math)


def add_inline(paragraph, text: str) -> None:
    parts = re.split(r"(\\\(.*?\\\)|\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("\\(") and part.endswith("\\)"):
            add_native_math(paragraph, part)
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        elif part.startswith("`") and part.endswith("`"):
            paragraph.add_run(part[1:-1]).font.name = "Consolas"
        else:
            paragraph.add_run(part)


def add_display_equation(document: Document, source: str) -> None:
    """Add a native Word Office-Math display equation from Markdown math."""
    equation = source.strip()
    if equation.startswith("\\["):
        equation = equation[2:]
    if equation.endswith("\\]"):
        equation = equation[:-2]

    paragraph = document.add_paragraph()
    paragraph.alignment = 1  # centre
    add_native_math(paragraph, equation)


def rebuild(markdown_path: Path, docx_path: Path) -> None:
    document = Document(docx_path) if docx_path.exists() else Document()
    clear_body(document)
    lines = markdown_path.read_text(encoding="utf-8").splitlines()
    index = 0
    while index < len(lines):
        line = lines[index]
        if not line.strip():
            index += 1
            continue
        if line.strip() == "\\[":
            equation_lines = [line]
            index += 1
            while index < len(lines):
                equation_lines.append(lines[index])
                if lines[index].strip() == "\\]":
                    index += 1
                    break
                index += 1
            add_display_equation(document, "\n".join(equation_lines))
            continue
        heading = re.match(r"^(#{1,6})\s+(.*)$", line)
        if heading:
            level = min(len(heading.group(1)), 3)
            document.add_heading(heading.group(2), level=level)
            index += 1
            continue
        if line.startswith("| ") and index + 1 < len(lines) and set(lines[index + 1]) <= {"|", "-", ":", " "}:
            rows = []
            while index < len(lines) and lines[index].startswith("|"):
                if set(lines[index]) <= {"|", "-", ":", " "}:
                    index += 1
                    continue
                rows.append([cell.strip() for cell in lines[index].strip("|").split("|")])
                index += 1
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Table Grid"
            for r, cells in enumerate(rows):
                for c, value in enumerate(cells):
                    table.cell(r, c).text = value
            continue
        bullet = re.match(r"^[-*]\s+(.*)$", line)
        if bullet:
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline(paragraph, bullet.group(1))
            index += 1
            continue
        paragraph = document.add_paragraph()
        add_inline(paragraph, line)
        index += 1
    document.save(docx_path)


def main() -> None:
    if len(sys.argv) == 3:
        rebuild(Path(sys.argv[1]), Path(sys.argv[2]))
        return
    if len(sys.argv) != 1:
        raise SystemExit("usage: build_thesis_chapters.py [input.md output.docx]")
    for number in range(1, 6):
        rebuild(THESIS / f"thesis_chapter_{number}.md", THESIS / f"CHPT{number}.docx")


if __name__ == "__main__":
    main()
